#!/usr/bin/env python3
"""Apply verified citation upgrades and the selected 13 insertions to the Word manuscript.

Does not invent venues. yang2025 is IJCAI 2025 (DOI 10.24963/ijcai.2025/1196), not 2024.
"""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = Path("/Users/jerryadamsfranklin/Downloads/Paper 1/Jerry_Manuscript_P1_Integrated.docx")
BAK = Path("/Users/jerryadamsfranklin/Downloads/Paper 1/Jerry_Manuscript_P1_Integrated.docx.bak")
OUT = SRC


def set_text(para: Paragraph, text: str) -> None:
    """Replace paragraph text, keeping the first run's formatting."""
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def insert_after(para: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = deepcopy(para._p)
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    if style:
        new_para.style = style
    set_text(new_para, text)
    return new_para


def insert_ref_after(para: Paragraph, text: str) -> Paragraph:
    return insert_after(para, text, style=para.style.name if para.style else "Body Text")


def main() -> None:
    if BAK.exists():
        shutil.copy2(BAK, SRC)
    else:
        shutil.copy2(SRC, BAK)
    doc = Document(str(SRC))
    p = doc.paragraphs

    # --- Section I: existing surveys (add Mao; Wang year 2025) ---
    set_text(
        p[17],
        "Relation to Existing Surveys. Several surveys address parameter-efficient fine-tuning or "
        "model compression individually. Critical reviews of PEFT methods (Xu et al., 2023) and "
        "broad methodological catalogues (Wang et al., 2025) cover adapter and low-rank families in "
        "depth but treat quantization and federated deployment as separate or secondary concerns. "
        "A dedicated LoRA-variant survey (Mao et al., 2025) organizes the rapidly growing adapter "
        "literature, including a federated-LoRA chapter, without connecting that taxonomy to "
        "quantization, memory systems, or a practitioner selection framework. Dedicated quantization "
        "surveys (Gong et al., 2024) review precision-reduction techniques thoroughly but do not "
        "connect them to adapter-based fine-tuning or to the practitioner method-selection problem. "
        "Surveys of federated low-rank adaptation (Yang et al., 2025) focus on the distributed "
        "protocol and presuppose, rather than survey, the single-device efficiency stack. This "
        "survey differs in three respects: it treats PEFT, quantization, memory optimization, and "
        "federated approaches as one connected design space rather than separate literatures; it "
        "centers a practitioner decision framework mapping hardware budgets to method choices; and "
        "it makes cross-method trade-offs explicit through a unified comparison rather than "
        "describing each family in isolation.",
    )

    # --- Section II: methodology ---
    set_text(
        p[24],
        "This survey follows a structured, protocol-driven search rather than a registered "
        "systematic review. Searches were conducted across arXiv, Semantic Scholar, OpenReview, "
        "OpenAlex, Crossref, and the IEEE Xplore metadata API, supplemented by manual searching of "
        "Google Scholar. Content published by the ACM was retrieved through OpenAlex and Crossref "
        "rather than through the ACM Digital Library interface directly. Boolean query strings were "
        "used where the source supports Boolean evaluation (arXiv, IEEE Xplore, Google Scholar); "
        "semantically equivalent keyword formulations were used for sources performing relevance "
        "matching (Semantic Scholar, OpenAlex, DBLP). Both formulations and the per-source logs are "
        "recorded in the companion repository. The coverage window is 1 January 2019 through "
        "30 June 2026.",
    )
    set_text(
        p[26],
        "Rather than attempting exhaustive coverage, we selected seminal papers introducing key "
        "techniques, empirical studies demonstrating practical trade-offs, and foundational "
        "theoretical work. The original selection comprised 42 core papers spanning Transformer "
        "architecture and optimization, parameter-efficient fine-tuning and quantization, memory "
        "optimization, and distributed training. A structured search of the 2024–2026 literature "
        "then added 13 in-scope methods and surveys, bringing the cited corpus to 55. For each "
        "paper, we extracted: method category and subcategory, theoretical mechanism, memory "
        "reduction ratio, accuracy retention on standard benchmarks, computational overhead, model "
        "sizes evaluated, and implementation availability. The taxonomy organizing this survey "
        "emerged through iterative analysis, identifying four primary categories based on the "
        "resource bottleneck each addresses.",
    )
    set_text(
        p[27],
        "To support reproduction of this selection, the exact query strings, the databases "
        "searched, and the inclusion and exclusion decisions are recorded in the companion "
        "repository accompanying this survey. Title-and-abstract screening of records published "
        "from January 2024 onward used a rule-encoded classifier calibrated to an author-resolved "
        "150-record sample; human verification for that band was applied to a citation- and "
        "section-weighted shortlist of includes rather than to every include. Records from 2019–2023 "
        "were subject to a targeted confirmation pass against the existing reference corpus, "
        "high-citation work, and core-method terms rather than an exhaustive screen. Records not "
        "examined under that bound are reported as a separate count in the companion repository and "
        "are not represented as screened. Where multiple papers reported the same technique, we "
        "retained the paper introducing the method and, where available, the most complete "
        "independent evaluation of it.",
    )

    # --- Section IV: LoRA variants ---
    set_text(
        p[53],
        "LoRA Variants. Subsequent research has systematically addressed different LoRA "
        "limitations. AdaLoRA (Zhang et al., 2023) recognizes that uniform rank across weight "
        "matrices is suboptimal and uses importance-based pruning, achieving 30–50% parameter "
        "reduction while maintaining accuracy. VeRA (Kopiczko et al., 2024) takes an opposing "
        "approach, sharing random matrices across layers and learning only per-layer scaling "
        "vectors for 90% parameter reduction at 1–2% accuracy cost. LoRA+ (Hayou et al., 2024) "
        "demonstrates that the A and B matrices play asymmetric roles; setting B's learning rate "
        "10–20× higher accelerates convergence by 1.5–2×, suggesting standard LoRA underutilizes "
        "its parameter budget due to optimization inefficiency. DoRA (Liu et al., 2024) explicitly "
        "decomposes updates into magnitude and directional components, closing 50–70% of the gap "
        "between LoRA and full fine-tuning, indicating the low-rank assumption itself is not the "
        "primary bottleneck. HydraLoRA (Tian et al., 2024) replaces LoRA's paired A/B with a shared "
        "A and multiple routed B matrices, which improves multi-domain adaptation without a "
        "hand-specified domain split but reintroduces a mixture-of-experts routing path that LoRA "
        "was designed to avoid. LoRA-GA (Wang, Yu, & Li, 2024) and PiSSA (Meng, Wang, & Zhang, 2024) "
        "change only initialization—gradient-aligned SVD in the former, principal components of W "
        "in the latter—so both converge faster than Gaussian/zero LoRA at the same rank without "
        "expanding the adapter architecture, and both pay an SVD cost at setup. LISA (Pan et al., "
        "2024) is not an adapter: it importance-samples layers and freezes most of the middle "
        "stack, matching LoRA's memory while updating full-rank weights in the sampled layers, so "
        "it does not yield a mergeable detachable module. These variants are not mutually "
        "exclusive: VeRA suits federated settings where communication dominates; DoRA provides "
        "superior accuracy when memory permits; LoRA+ offers immediate speedups (Table 2).",
    )

    # --- Section IV: QLoRA / IR-QLoRA ---
    set_text(
        p[59],
        "The complete memory savings are dramatic. For a 7B model, full fine-tuning requires 132GB; "
        "standard LoRA with FP16 reduces this to approximately 29GB; QLoRA achieves 5–6GB with "
        "paged optimizers, a 22× reduction enabling consumer GPU training. Training speed decreases "
        "by 30–40% due to on-the-fly dequantization, and quality typically remains within 1–2% of "
        "full-precision LoRA on standard benchmarks, though precision-sensitive tasks may show "
        "larger gaps. IR-QLoRA (Qin et al., 2024) reduces that residual drop by calibrating the "
        "quantizer to retain more information and adding an elastic connection around LoRA, but it "
        "remains a LoRA-on-quantized-base recipe rather than a jointly learned quantizer.",
    )

    # --- Section V: PTQ vs QAT ---
    set_text(
        p[73],
        "Selecting between PTQ and QAT. The choice between post-training quantization and "
        "quantization-aware training is driven by three factors: the availability of a training "
        "budget, the target bit-width, and tolerance for accuracy loss. Post-training quantization "
        "(GPTQ, AWQ) requires no retraining and suits models that must be compressed quickly or "
        "updated frequently, with AWQ preferred when calibration speed matters and GPTQ when "
        "maximum compression quality at a one-time cost is acceptable. SpinQuant (Liu et al., 2025) "
        "extends this PTQ family by learning rotation matrices that absorb activation outliers "
        "before quantization, narrowing the 4-bit gap to full precision without a QAT run; it is "
        "still a calibration procedure, not a training-time method, and it does not change the "
        "fine-tuning recipe. Quantization-aware training becomes justified only when the target "
        "precision is aggressive enough, typically at or below 3-bit, that post-training methods "
        "incur unacceptable degradation, because it recovers accuracy at the cost of a full "
        "training run. EfficientQAT (Chen et al., 2025) weakens that cost objection by training "
        "block-wise and then freezing weights while tuning only step sizes, making 2-bit QAT of "
        "70B-class models feasible on a single 80GB GPU—but that remains far more expensive than a "
        "GPTQ, AWQ, or SpinQuant calibration pass. Above 4-bit, the accuracy advantage of QAT is "
        "usually too small to justify its cost, and PTQ dominates on a cost-quality basis.",
    )

    # --- Section VI: FlashAttention-2 year ---
    set_text(
        p[80],
        "Self-attention's quadratic memory complexity with respect to sequence length has been a "
        "fundamental bottleneck. FlashAttention (Dao et al., 2022) redesigns attention computation "
        "to process in blocks that fit in GPU on-chip SRAM, exploiting the mathematical property "
        "that softmax can be computed incrementally. The full N×N attention matrix is never "
        "materialized, reducing memory from O(N²) to O(N). Counterintuitively, this also achieves "
        "2–4× speedup because reducing high-bandwidth memory accesses more than compensates for "
        "additional arithmetic on memory-bandwidth-limited GPUs. FlashAttention-2 (Dao, 2024) "
        "extends these gains through improved parallelization, reaching 1.5–2× additional speedup. "
        "The technique is nearly universally beneficial and should be enabled by default on "
        "compatible hardware.",
    )

    # --- Section VI: new 6.4 GaLore / AdaRankGrad, renumber combining ---
    set_text(p[83], "6.5 Combining Memory Optimizations")
    galore = (
        "Gradient-projection methods attack optimizer-state memory rather than activations. "
        "GaLore (Zhao et al., 2024) stores Adam moments in a low-rank subspace of the gradient, "
        "enabling full-parameter training of a 7B model on 24GB hardware without freezing the base "
        "weights; unlike LoRA it does not produce a detachable adapter, the projection must be "
        "refreshed by periodic SVD, and the strongest published results are for pretraining rather "
        "than downstream fine-tuning. AdaRankGrad (Refael et al., 2025) makes that rank adaptive as "
        "training proceeds, avoiding a fixed GaLore rank, at the cost of extra subspace estimation "
        "whenever the rank changes."
    )
    heading_style = p[83].style
    body_style = p[84].style
    insert_after(p[82], galore, style=body_style)
    insert_after(p[82], "6.4 Gradient Low-Rank Projection", style=heading_style)

    # --- Section VII: federated PEFT (reattribute FlexLoRA to Bai) ---
    fed = (
        "FlexLoRA (Bai et al., 2024)—a different method from the FlexLoRA preprint of "
        "Bayati et al. (2023)—lets clients train at different LoRA ranks and reconstructs "
        "a full-size update by SVD, so well-resourced clients are no longer capped by the "
        "weakest participant; the redistribution assumes the low-rank factors remain "
        "comparable across rounds, which is fragile when tasks differ sharply. FFA-LoRA "
        "(Sun et al., 2024) freezes the random A matrix and trains only B under "
        "differential privacy, halving communication and reducing DP-noise amplification, "
        "at the price of a strictly smaller hypothesis class than vanilla LoRA. OpenFedLLM "
        "(Ye et al., 2024) is a federation-of-LoRA codebase rather than a new adapter; its "
        "result that federated instruction tuning can beat local training is demonstrated "
        "mainly on instruction-following and alignment, not on arbitrary downstream heads. "
        "FLoRA (Zhang et al., 2024) adjusts ranks dynamically based on observed client "
        "contribution quality. QLoRA's 5–6GB requirement further expands participation to "
        "consumer GPUs and gaming laptops, with each client independently choosing "
        "quantization strategy while the server aggregates in the common LoRA parameter space."
    )
    for para in doc.paragraphs:
        if "FlexLoRA (Bayati et al., 2023)" in para.text:
            set_text(para, fed)
            break
    else:
        raise SystemExit("federated FlexLoRA paragraph not found")

    for para in doc.paragraphs:
        if "resource-adaptive rank allocation" in para.text:
            set_text(
                para,
                para.text.replace(
                    "federated learning with resource-adaptive rank allocation enables",
                    "federated learning with resource-adaptive rank allocation (Bai et al., 2024) enables",
                ),
            )
            break

    # --- Section 10.4 limitations ---
    for para in doc.paragraphs:
        if para.text.startswith("Third, the selection is curated rather than exhaustive"):
            set_text(
                para,
                "Third, the selection is curated rather than exhaustive. We prioritized seminal "
                "methods, widely adopted implementations, and studies reporting sufficient detail "
                "to extract efficiency measures, which biases coverage toward methods that achieved "
                "community traction. Techniques that are promising but not yet widely evaluated may "
                "be underrepresented. Title-and-abstract screening of the 2024–2026 band was "
                "automated under an encoded inclusion rule and then human-verified on a ranked "
                "shortlist, not on every include; it is not dual independent screening, and it is "
                "not a registered systematic review.",
            )
        if para.text.startswith("Finally, the literature in this area moves quickly"):
            set_text(
                para,
                "Finally, the literature in this area moves quickly. The search covers publications "
                "through June 2026, and methods introduced after that point are not represented. "
                "Given the pace of development in parameter-efficient fine-tuning and quantization, "
                "some specific numerical claims will date faster than the structural trade-offs the "
                "survey describes; the decision framework in Section IX is intended to remain "
                "applicable even as individual method rankings shift.",
            )

    # --- References: upgrades + new entries, keep alphabetical ---
    ref_map = {
        "Dao, T. (2023). FlashAttention-2":
            "Dao, T. (2024). FlashAttention-2: Faster attention with better parallelism and work partitioning. Proceedings of the International Conference on Learning Representations (ICLR).",
        "Wang, L., Chen, S., Jiang, L., Pan, S., Cai, R., Yang, S., & Yang, F. (2024). Parameter-efficient fine-tuning in large models":
            "Wang, L., Chen, S., Jiang, L., Pan, S., Cai, R., Yang, S., & Yang, F. (2025). Parameter-efficient fine-tuning in large language models: A survey of methodologies. Artificial Intelligence Review, 58, 227. https://doi.org/10.1007/s10462-025-11236-4",
        "Yang, Y., Long, G., Lu, Q., Zhu, L., Jiang, J., & Zhang, C. (2025). Federated low-rank adaptation":
            "Yang, Y., Long, G., Lu, Q., Zhu, L., Jiang, J., & Zhang, C. (2025). Federated low-rank adaptation for foundation models: A survey. Proceedings of the Thirty-Fourth International Joint Conference on Artificial Intelligence (IJCAI), 10779–10787. https://doi.org/10.24963/ijcai.2025/1196",
    }
    for para in p:
        t = para.text
        for prefix, replacement in ref_map.items():
            if t.startswith(prefix):
                set_text(para, replacement)

    def find_para(prefix: str) -> Paragraph:
        for para in doc.paragraphs:
            if para.text.startswith(prefix):
                return para
        raise SystemExit(f"no paragraph starting {prefix!r}")

    def insert_before(anchor_prefix: str, ref_text: str) -> None:
        para = find_para(anchor_prefix)
        prev = para._p.getprevious()
        while prev is not None and prev.tag != qn("w:p"):
            prev = prev.getprevious()
        if prev is None:
            raise SystemExit(f"no previous paragraph before {anchor_prefix!r}")
        insert_after(Paragraph(prev, para._parent), ref_text, style=para.style)

    def insert_after_prefix(prefix: str, ref_text: str) -> None:
        para = find_para(prefix)
        insert_after(para, ref_text, style=para.style)

    insert_before("Bayati, B.",
        "Bai, J., Chen, D., Qian, B., Yao, L., & Li, Y. (2024). Federated fine-tuning of large language models under heterogeneous tasks and client resources. Advances in Neural Information Processing Systems, 37.")
    insert_before("Chen, T., Xu, B.",
        "Chen, M., Shao, W., Xu, P., Wang, J., Gao, P., Zhang, K., & Luo, P. (2025). EfficientQAT: Efficient quantization-aware training for large language models. Proceedings of the 63rd Annual Meeting of the Association for Computational Linguistics (Volume 1: Long Papers), 10081–10100. https://doi.org/10.18653/v1/2025.acl-long.498")
    insert_before("Loshchilov, I.",
        "Liu, Z., Zhao, C., Fedorov, I., Soran, B., Choudhary, D., Krishnamoorthi, R., Chandra, V., Tian, Y., & Blankevoort, T. (2025). SpinQuant: LLM quantization with learned rotations. Proceedings of the International Conference on Learning Representations (ICLR).")
    insert_before("McMahan, B.",
        "Mao, Y., Ge, Y., Fan, Y., Xu, W., Mi, Y., Hu, Z., & Gao, Y. (2025). A survey on LoRA of large language models. Frontiers of Computer Science, 19, 197605. https://doi.org/10.1007/s11704-024-40663-9")
    insert_before("Micikevicius, P.",
        "Meng, F., Wang, Z., & Zhang, M. (2024). PiSSA: Principal singular values and singular vectors adaptation of large language models. Advances in Neural Information Processing Systems, 37.")
    insert_before("Pfeiffer, J.",
        "Pan, R., Liu, X., Diao, S., Pi, R., Zhang, J., Han, C., & Zhang, T. (2024). LISA: Layerwise importance sampling for memory-efficient large language model fine-tuning. Advances in Neural Information Processing Systems, 37.")
    insert_before("Rafailov, R.",
        "Qin, H., Ma, X., Zheng, X., Li, X., Zhang, Y., Liu, S., Luo, J., Liu, X., & Magno, M. (2024). Accurate LoRA-finetuning quantization of LLMs via information retention. Proceedings of the 41st International Conference on Machine Learning.")
    insert_before("Ren, J.",
        "Refael, Y., Svirsky, J., Shustin, B., Huleihel, W., & Lindenbaum, O. (2025). AdaRankGrad: Adaptive gradient-rank and moments for memory-efficient LLMs training and fine-tuning. Proceedings of the International Conference on Learning Representations (ICLR).")
    insert_before("Vaswani, A.",
        "Sun, Y., Li, Z., Li, Y., & Ding, B. (2024). Improving LoRA in privacy-preserving federated learning. Proceedings of the International Conference on Learning Representations (ICLR).")
    insert_before("Vaswani, A.",
        "Tian, C., Shi, Z., Guo, Z., Li, L., & Xu, C.-Z. (2024). HydraLoRA: An asymmetric LoRA architecture for efficient fine-tuning. Advances in Neural Information Processing Systems, 37.")
    insert_before("Xiao, G.",
        "Wang, S., Yu, L., & Li, J. (2024). LoRA-GA: Low-rank adaptation with gradient approximation. Advances in Neural Information Processing Systems, 37.")
    insert_before("Zaken, E. B.",
        "Ye, R., Wang, W., Chai, J., Li, D., Li, Z., Xu, Y., Du, Y., Wang, Y., & Chen, S. (2024). OpenFedLLM: Training large language models on decentralized private data via federated learning. Proceedings of the 30th ACM SIGKDD Conference on Knowledge Discovery and Data Mining, 6137–6147. https://doi.org/10.1145/3637528.3671582")
    insert_after_prefix("Zhang, Q., Chen, M.",
        "Zhao, J., Zhang, Z., Chen, B., Wang, Z., Anandkumar, A., & Tian, Y. (2024). GaLore: Memory-efficient LLM training by gradient low-rank projection. Proceedings of the 41st International Conference on Machine Learning.")

    doc.save(str(OUT))
    print(f"wrote {OUT}")
    print(f"backup {BAK}")


if __name__ == "__main__":
    main()
